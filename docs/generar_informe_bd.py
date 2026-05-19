"""Genera el archivo Word del Informe de Base de Datos.

Requiere: pip install python-docx

Uso:
    python docs/generar_informe_bd.py

Produce:
    docs/Informe_Base_Datos.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.dirname(os.path.abspath(__file__))
DIAGRAMS = os.path.join(DOCS, "diagrams")

# Paleta corporativa
COLOR_NAVY    = RGBColor(0x00, 0x3A, 0x6B)
COLOR_BLUE    = RGBColor(0x00, 0x5C, 0xAB)
COLOR_BLUE_LT = RGBColor(0x01, 0x83, 0xEF)
COLOR_ORANGE  = RGBColor(0xFF, 0x94, 0x00)
COLOR_GREEN   = RGBColor(0x00, 0xA5, 0x0B)
COLOR_GREY    = RGBColor(0x5B, 0x6B, 0x7A)
COLOR_CODE_BG = "F1F5F9"


# =========================================================
# Helpers
# =========================================================
def _set_default_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    for level, size, color, sb, sa in [
        ("Heading 1", 20, COLOR_NAVY, Pt(16), Pt(6)),
        ("Heading 2", 16, COLOR_BLUE, Pt(14), Pt(4)),
        ("Heading 3", 13, COLOR_BLUE, Pt(10), Pt(3)),
        ("Heading 4", 11, COLOR_NAVY, Pt(8),  Pt(2)),
    ]:
        s = doc.styles[level]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True
        s.paragraph_format.space_before = sb
        s.paragraph_format.space_after = sa


def _add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def _add_p(doc, text, **kwargs):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if kwargs.get("bold"):
        run.bold = True
    if kwargs.get("italic"):
        run.italic = True
    if kwargs.get("size"):
        run.font.size = Pt(kwargs["size"])
    if kwargs.get("color"):
        run.font.color.rgb = kwargs["color"]
    if kwargs.get("align"):
        p.alignment = kwargs["align"]
    return p


def _add_bullet_list(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            run = p.add_run(it[0])
            run.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def _add_numbered_list(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(it)


def _shade(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _add_table(doc, headers, rows, header_fill="005CAB"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        _shade(hdr[i], header_fill)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for c, value in enumerate(row):
            cells[c].text = ""
            run = cells[c].paragraphs[0].add_run(str(value))
            run.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _add_code(doc, code: str):
    """Inserta un bloque de código (SQL/Python) con fondo gris claro."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _shade(cell, COLOR_CODE_BG)
    # Limpia el párrafo por defecto y agrega líneas
    cell.text = ""
    p = cell.paragraphs[0]
    for line in code.splitlines():
        if p.text != "" or any(r.text for r in p.runs):
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_NAVY
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _add_image(doc, path, caption=None, width_cm=15.5):
    if not os.path.isfile(path):
        _add_p(doc, f"[imagen no encontrada: {path}]", italic=True, color=COLOR_GREY)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cap.add_run(caption)
        crun.italic = True
        crun.font.size = Pt(9)
        crun.font.color.rgb = COLOR_GREY


def _add_footer(doc, text):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = COLOR_GREY


def _add_portada(doc, titulo, subtitulo):
    sect = doc.sections[0]
    sect.top_margin = Cm(2.5); sect.bottom_margin = Cm(2.5)
    sect.left_margin = Cm(2.5); sect.right_margin = Cm(2.5)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDAD DEL MAGDALENA")
    run.font.size = Pt(14); run.bold = True; run.font.color.rgb = COLOR_BLUE

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Control Estadístico de Procesos — 2026-1")
    run.font.size = Pt(11); run.italic = True; run.font.color.rgb = COLOR_GREY

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titulo)
    run.font.size = Pt(26); run.bold = True; run.font.color.rgb = COLOR_NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitulo)
    run.font.size = Pt(15); run.italic = True; run.font.color.rgb = COLOR_BLUE

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sistema de Control Estadístico de Calidad")
    run.font.size = Pt(12); run.bold = True; run.font.color.rgb = COLOR_ORANGE

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Frutas, hortalizas y plantas medicinales")
    run.font.size = Pt(11); run.font.color.rgb = COLOR_GREY

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Santa Marta, mayo de 2026")
    run.font.size = Pt(11); run.font.color.rgb = COLOR_GREY

    doc.add_page_break()


def _add_toc(doc, items):
    _add_heading(doc, "Tabla de contenidos", level=1)
    for txt, level in items:
        p = doc.add_paragraph()
        if level > 1:
            p.paragraph_format.left_indent = Cm((level - 1) * 0.5)
        run = p.add_run(txt)
        run.font.size = Pt(11)
        if level == 1:
            run.bold = True; run.font.color.rgb = COLOR_NAVY
        else:
            run.font.color.rgb = COLOR_GREY
    doc.add_page_break()


# =========================================================
# INFORME DE BASE DE DATOS
# =========================================================
def generar():
    doc = Document()
    _set_default_styles(doc)
    _add_portada(doc, "Informe de Base de Datos",
                 "Modelo, esquema, configuración y operación")
    _add_footer(doc, "Universidad del Magdalena · Sistema CEC · Informe de Base de Datos")

    _add_toc(doc, [
        ("1. Introducción", 1),
        ("2. Elección tecnológica", 1),
        ("3. Arquitectura de conexión", 1),
        ("4. Modelo entidad-relación", 1),
        ("5. Esquema SQL (DDL)", 1),
        ("6. Detalle de las tablas", 1),
        ("7. Tipos de datos y justificación", 1),
        ("8. Restricciones e índices", 1),
        ("9. Configuración y variables de entorno", 1),
        ("10. Operaciones CRUD soportadas", 1),
        ("11. Patrones de acceso desde la aplicación", 1),
        ("12. Particularidades para entornos serverless", 1),
        ("13. Seguridad", 1),
        ("14. Backups, mantenimiento y restauración", 1),
        ("15. Consultas útiles", 1),
        ("16. Migración y portabilidad", 1),
        ("17. Métricas y monitoreo", 1),
        ("18. Limitaciones conocidas", 1),
        ("19. Glosario", 1),
    ])

    # 1
    _add_heading(doc, "1. Introducción", 1)
    doc.add_paragraph(
        "Este documento describe el diseño y operación de la base de datos que da "
        "soporte al Sistema de Control Estadístico de Calidad. El objetivo es servir "
        "como referencia técnica para entender qué información se persiste, cómo "
        "se organiza, cómo se accede a ella desde la aplicación y qué consideraciones "
        "deben tenerse en cuenta para su mantenimiento y evolución."
    )
    doc.add_paragraph(
        "El sistema utiliza una arquitectura de doble backend con resolución automática: "
        "PostgreSQL (en Supabase) en producción y SQLite en desarrollo local. Ambos "
        "backends comparten la misma interfaz pública en Python, lo que permite alternar "
        "entre ellos simplemente configurando o no la variable de entorno "
        "DATABASE_URL."
    )

    # 2
    _add_heading(doc, "2. Elección tecnológica", 1)
    _add_heading(doc, "2.1 PostgreSQL en producción", 2)
    doc.add_paragraph(
        "PostgreSQL es el motor relacional de referencia para aplicaciones modernas: "
        "ACID-compliant, con soporte amplio de tipos (incluyendo JSONB), índices "
        "avanzados, full-text search, replicación y un ecosistema maduro de herramientas. "
        "La versión que provee Supabase es PostgreSQL 15+."
    )
    _add_heading(doc, "2.2 Supabase como proveedor", 2)
    _add_bullet_list(doc, [
        ("Plan gratuito generoso: ", "500 MB de almacenamiento, 2 GB de transferencia, "
         "ilimitadas APIs y filas. Suficiente para el prototipo académico."),
        ("Administración cero: ", "Supabase se encarga de parches, backups, alta "
         "disponibilidad y monitoreo."),
        ("API REST y RPC automáticas: ", "PostgREST expone las tablas como endpoints "
         "(no usado en este proyecto, pero disponible para evoluciones)."),
        ("Auth integrada: ", "Supabase Auth permite agregar usuarios y RLS sin "
         "infraestructura adicional."),
        ("Edge functions: ", "ejecución de lógica próxima al cliente para casos "
         "específicos."),
    ])

    _add_heading(doc, "2.3 SQLite como fallback local", 2)
    doc.add_paragraph(
        "Para el desarrollo y las pruebas se mantiene un backend SQLite que escribe en "
        "data/control_calidad.db. Esto evita la dependencia de una conexión a Supabase "
        "durante el desarrollo, permite trabajar offline y simplifica las pruebas "
        "automáticas. SQLite NO es apto para Vercel porque el filesystem en serverless "
        "es efímero."
    )

    _add_heading(doc, "2.4 URL del proyecto Supabase", 2)
    doc.add_paragraph("URL del proyecto: https://kicalhpqppkknqtjhtml.supabase.co")

    # 3
    _add_heading(doc, "3. Arquitectura de conexión", 1)
    _add_image(doc, os.path.join(DIAGRAMS, "arquitectura.png"),
               caption="Figura 1. Arquitectura general de la solución. La capa de "
               "persistencia (Supabase) recibe conexiones del backend Flask vía el "
               "Transaction Pooler.",
               width_cm=16)
    _add_heading(doc, "3.1 Componentes", 2)
    _add_bullet_list(doc, [
        ("Cliente (navegador): ", "se comunica únicamente con el backend Flask vía HTTPS. "
         "Nunca habla directamente con la base de datos."),
        ("Backend (Vercel Serverless): ", "función Python que ejecuta Flask. Cada "
         "petición HTTP genera una nueva invocación que abre y cierra su conexión."),
        ("Transaction Pooler de Supabase (PgBouncer): ", "puerto 6543, IPv4. Reutiliza "
         "conexiones físicas a Postgres entre invocaciones; indispensable para "
         "serverless."),
        ("PostgreSQL: ", "el motor donde residen las tablas estudios y muestras."),
    ])

    _add_heading(doc, "3.2 Opciones de conexión a Supabase", 2)
    _add_table(doc, ["Modo", "Puerto", "Host", "Protocolo", "Apto en Vercel"], [
        ("Direct connection", "5432", "db.<ref>.supabase.co", "IPv6", "No (Vercel no tiene IPv6 saliente)"),
        ("Session pooler",    "5432", "aws-0-<region>.pooler.supabase.com", "IPv4", "Sí (alternativa)"),
        ("Transaction pooler","6543", "aws-0-<region>.pooler.supabase.com", "IPv4", "Sí (recomendada)"),
    ])
    doc.add_paragraph(
        "En este proyecto se usa el Transaction Pooler porque maneja eficientemente las "
        "conexiones efímeras de las funciones serverless."
    )

    # 4
    _add_heading(doc, "4. Modelo entidad-relación", 1)
    _add_image(doc, os.path.join(DIAGRAMS, "modelo_datos.png"),
               caption="Figura 2. Modelo entidad-relación. Una fila en estudios "
               "agrupa varios subgrupos en muestras.",
               width_cm=15)
    doc.add_paragraph(
        "El modelo es deliberadamente simple — solo dos tablas. La granularidad de "
        "análisis del CEP es el subgrupo, no la medición individual; por eso las "
        "mediciones de un mismo subgrupo se guardan juntas en un campo JSONB en "
        "lugar de normalizarse en una tercera tabla."
    )

    # 5
    _add_heading(doc, "5. Esquema SQL (DDL)", 1)
    doc.add_paragraph("DDL completo aplicado en Supabase (archivo supabase/schema.sql):")
    _add_code(doc, """CREATE TABLE IF NOT EXISTS public.estudios (
    id              BIGSERIAL PRIMARY KEY,
    nombre          TEXT NOT NULL,
    producto        TEXT NOT NULL,
    tipo            TEXT NOT NULL CHECK (tipo IN ('variable', 'atributo')),
    caracteristica  TEXT NOT NULL,
    unidad          TEXT,
    analista        TEXT,
    lote            TEXT,
    tipo_grafico    TEXT NOT NULL
                    CHECK (tipo_grafico IN ('xr', 'xs', 'p', 'np', 'c', 'u')),
    lsl             DOUBLE PRECISION,
    usl             DOUBLE PRECISION,
    tamano_subgrupo INTEGER
                    CHECK (tamano_subgrupo IS NULL OR
                           (tamano_subgrupo BETWEEN 2 AND 25)),
    notas           TEXT,
    fecha_creacion  TIMESTAMPTZ NOT NULL DEFAULT NOW()
);

CREATE TABLE IF NOT EXISTS public.muestras (
    id            BIGSERIAL PRIMARY KEY,
    estudio_id    BIGINT NOT NULL
                  REFERENCES public.estudios(id) ON DELETE CASCADE,
    subgrupo      INTEGER NOT NULL CHECK (subgrupo >= 1),
    valores       JSONB NOT NULL,
    fecha_muestra TIMESTAMPTZ NOT NULL DEFAULT NOW()
);

CREATE INDEX IF NOT EXISTS idx_muestras_estudio
    ON public.muestras(estudio_id);
CREATE INDEX IF NOT EXISTS idx_estudios_fecha
    ON public.estudios(fecha_creacion DESC);
CREATE INDEX IF NOT EXISTS idx_estudios_producto
    ON public.estudios(producto);""")

    doc.add_paragraph(
        "Además, se incluye una vista de conveniencia para el dashboard:"
    )
    _add_code(doc, """CREATE OR REPLACE VIEW public.v_estudios_resumen AS
SELECT
    e.id, e.nombre, e.producto, e.tipo, e.caracteristica,
    e.tipo_grafico, e.analista, e.lote, e.fecha_creacion,
    COUNT(m.id) AS n_muestras
FROM public.estudios e
LEFT JOIN public.muestras m ON m.estudio_id = e.id
GROUP BY e.id
ORDER BY e.fecha_creacion DESC;""")

    # 6
    _add_heading(doc, "6. Detalle de las tablas", 1)
    _add_heading(doc, "6.1 Tabla estudios", 2)
    doc.add_paragraph("Almacena la trazabilidad y configuración de cada estudio de control:")
    _add_table(doc, ["Columna", "Tipo", "Obligatorio", "Descripción"], [
        ("id",                "BIGSERIAL",        "Sí (auto)",
         "Clave primaria autoincremental"),
        ("nombre",            "TEXT",             "Sí",
         "Nombre descriptivo del estudio"),
        ("producto",          "TEXT",             "Sí",
         "Producto bajo control (mango, sábila…)"),
        ("tipo",              "TEXT",             "Sí",
         "'variable' o 'atributo'"),
        ("caracteristica",    "TEXT",             "Sí",
         "Característica controlada (peso, manchas…)"),
        ("unidad",            "TEXT",             "No",
         "Unidad de medida (g, cm, °Bx, %)"),
        ("analista",          "TEXT",             "No",
         "Nombre del analista responsable"),
        ("lote",              "TEXT",             "No",
         "Identificador del lote / cosecha"),
        ("tipo_grafico",      "TEXT",             "Sí",
         "xr | xs | p | np | c | u"),
        ("lsl",               "DOUBLE PRECISION", "No",
         "Límite inferior de especificación"),
        ("usl",               "DOUBLE PRECISION", "No",
         "Límite superior de especificación"),
        ("tamano_subgrupo",   "INTEGER",          "X̄-R / X̄-S",
         "Mediciones por subgrupo (2..25)"),
        ("notas",             "TEXT",             "No",
         "Observaciones del estudio"),
        ("fecha_creacion",    "TIMESTAMPTZ",      "Auto (NOW)",
         "Fecha y hora UTC"),
    ])

    _add_heading(doc, "6.2 Tabla muestras", 2)
    doc.add_paragraph(
        "Almacena cada subgrupo como una fila. Las mediciones del subgrupo se "
        "guardan en una columna JSONB para preservar la integridad atómica del "
        "subgrupo y simplificar las consultas."
    )
    _add_table(doc, ["Columna", "Tipo", "Obligatorio", "Descripción"], [
        ("id",            "BIGSERIAL",     "Sí (auto)", "Clave primaria autoincremental"),
        ("estudio_id",    "BIGINT",        "Sí",        "FK a estudios.id (ON DELETE CASCADE)"),
        ("subgrupo",      "INTEGER",       "Sí",        "Número de subgrupo (1..N)"),
        ("valores",       "JSONB",         "Sí",        "Array de mediciones o [defectivos, n]"),
        ("fecha_muestra", "TIMESTAMPTZ",   "Auto (NOW)","Fecha de captura"),
    ])

    _add_heading(doc, "6.3 Formato del campo 'valores'", 2)
    doc.add_paragraph("El contenido depende del tipo_grafico del estudio asociado:")
    _add_table(doc, ["Tipo de gráfico", "Formato JSON", "Ejemplo"], [
        ("xr / xs", "[v1, v2, ..., vn]", "[248.5, 251.2, 249.8, 250.1, 249.9]"),
        ("p",       "[defectivos, n]",   "[7, 100]"),
        ("np",      "[defectivos, n]",   "[4, 50]"),
        ("c",       "[defectos]",        "[6]"),
        ("u",       "[defectos, n]",     "[12, 8]"),
    ])

    # 7
    _add_heading(doc, "7. Tipos de datos y justificación", 1)
    _add_bullet_list(doc, [
        ("BIGSERIAL: ", "PK autoincremental de 64 bits. Cabida holgada para crecimiento."),
        ("TEXT: ", "preferido sobre VARCHAR(n) en Postgres porque no impone límite "
         "artificial y el costo es idéntico."),
        ("DOUBLE PRECISION: ", "para LSL/USL — las especificaciones técnicas pueden "
         "requerir decimales (p. ej. pH = 6.85)."),
        ("INTEGER: ", "para tamaño de subgrupo y número de subgrupo."),
        ("TIMESTAMPTZ: ", "siempre con zona horaria. Se almacena en UTC; las "
         "aplicaciones convierten al mostrar."),
        ("JSONB: ", "binario, indexable, eficiente. Permite almacenar la unidad "
         "lógica del subgrupo sin perder consultabilidad."),
    ])

    # 8
    _add_heading(doc, "8. Restricciones e índices", 1)
    _add_heading(doc, "8.1 CHECK constraints", 2)
    _add_table(doc, ["Tabla", "Columna", "Regla", "Por qué"], [
        ("estudios", "tipo", "IN ('variable', 'atributo')", "Domina los valores válidos"),
        ("estudios", "tipo_grafico", "IN (xr, xs, p, np, c, u)",
         "Solo los 6 gráficos implementados"),
        ("estudios", "tamano_subgrupo", "NULL o 2..25",
         "Constantes de Shewhart tabuladas hasta n=25"),
        ("muestras", "subgrupo", ">= 1",
         "Numeración natural"),
    ])

    _add_heading(doc, "8.2 Foreign key", 2)
    doc.add_paragraph(
        "muestras.estudio_id REFERENCES estudios(id) ON DELETE CASCADE. Al eliminar "
        "un estudio se borran automáticamente todas sus muestras, evitando filas "
        "huérfanas."
    )

    _add_heading(doc, "8.3 Índices", 2)
    _add_table(doc, ["Nombre", "Tabla", "Columna(s)", "Propósito"], [
        ("idx_muestras_estudio", "muestras", "estudio_id",
         "JOIN frecuente al cargar la página de detalle"),
        ("idx_estudios_fecha", "estudios", "fecha_creacion DESC",
         "Listado por más recientes (dashboard)"),
        ("idx_estudios_producto", "estudios", "producto",
         "Filtro/búsqueda por producto"),
    ])

    # 9
    _add_heading(doc, "9. Configuración y variables de entorno", 1)
    doc.add_paragraph("La aplicación elige el backend según el entorno:")
    _add_table(doc, ["Variable", "Obligatoria", "Descripción"], [
        ("DATABASE_URL", "Sí en producción",
         "URI de Postgres (Transaction Pooler de Supabase)"),
        ("SECRET_KEY", "Recomendada", "Clave de sesión Flask"),
        ("DB_PATH", "No", "Override de ruta SQLite local"),
        ("VERCEL", "Auto", "Vercel la define; se usa para SQLite en /tmp"),
    ])

    _add_heading(doc, "9.1 Formato de DATABASE_URL", 2)
    _add_code(doc, """postgresql://postgres.kicalhpqppkknqtjhtml:[YOUR-PASSWORD]@aws-0-<region>.pooler.supabase.com:6543/postgres""")
    doc.add_paragraph(
        "Donde [YOUR-PASSWORD] se reemplaza por la contraseña real (obtenida de "
        "Supabase Project Settings → Database). La región (us-east-1, sa-east-1, etc.) "
        "depende de dónde se creó el proyecto."
    )

    # 10
    _add_heading(doc, "10. Operaciones CRUD soportadas", 1)
    doc.add_paragraph(
        "La interfaz pública de la capa de datos está en app/database.py (dispatcher). "
        "Las funciones disponibles son:"
    )
    _add_table(doc, ["Función", "Descripción"], [
        ("init_db()", "Crea las tablas si no existen (idempotente)"),
        ("crear_estudio(payload)", "Inserta un estudio y devuelve su id"),
        ("listar_estudios()", "Devuelve todos los estudios ordenados por fecha desc"),
        ("obtener_estudio(id)", "Devuelve un estudio por id"),
        ("eliminar_estudio(id)", "Elimina estudio + sus muestras (cascada)"),
        ("agregar_muestra(estudio_id, subgrupo, valores)", "Una muestra individual"),
        ("agregar_muestras_bulk(estudio_id, muestras)", "Inserción masiva eficiente"),
        ("listar_muestras(estudio_id)", "Devuelve las muestras de un estudio"),
        ("eliminar_muestras(estudio_id)", "Borra todas las muestras de un estudio"),
    ])

    # 11
    _add_heading(doc, "11. Patrones de acceso desde la aplicación", 1)
    _add_heading(doc, "11.1 Flujo de creación de estudio", 2)
    _add_numbered_list(doc, [
        "POST /api/estudios recibe el JSON con la trazabilidad y opcionalmente las muestras.",
        "crear_estudio(payload) inserta en estudios y devuelve el id.",
        "Si vienen muestras, agregar_muestras_bulk las inserta en una sola transacción.",
        "El cliente recibe el id y redirige a /estudio/<id>.",
    ])

    _add_heading(doc, "11.2 Flujo de carga por Excel", 2)
    _add_numbered_list(doc, [
        "POST /api/estudios/upload recibe el archivo .xlsx (multipart).",
        "parse_excel valida estructura y produce el payload.",
        "Se llaman las mismas funciones de la BD (crear_estudio + bulk).",
        "Se redirige a la página de análisis del estudio creado.",
    ])

    _add_heading(doc, "11.3 Flujo de lectura para análisis", 2)
    _add_numbered_list(doc, [
        "GET /estudio/<id> renderiza el HTML.",
        "obtener_estudio(id) + listar_muestras(id) cargan los datos.",
        "El cliente solicita los cálculos al backend con POST /api/analisis/*.",
        "Los algoritmos NumPy/SciPy ejecutan sobre los datos en memoria y devuelven JSON.",
    ])

    # 12
    _add_heading(doc, "12. Particularidades para entornos serverless", 1)
    _add_heading(doc, "12.1 Una conexión por petición", 2)
    doc.add_paragraph(
        "Cada invocación Lambda en Vercel es independiente y efímera. La estrategia "
        "elegida es abrir y cerrar la conexión a Postgres en cada petición. Esto es "
        "ineficiente sin un pooler, pero con el Transaction Pooler de Supabase se "
        "vuelve perfectamente viable."
    )

    _add_heading(doc, "12.2 Prepared statements desactivados", 2)
    doc.add_paragraph(
        "psycopg 3 por defecto crea prepared statements para optimizar consultas "
        "repetidas. En el Transaction Pooler (PgBouncer en modo transaction), las "
        "conexiones físicas se reutilizan entre clientes, lo que rompe la "
        "correspondencia entre PREPARE/EXECUTE y produce el error 'prepared "
        "statement \"_pg3_0\" already exists'. La solución implementada es:"
    )
    _add_code(doc, """conn = psycopg.connect(
    DATABASE_URL,
    row_factory=dict_row,
    autocommit=False,
    prepare_threshold=None,  # desactiva preparación automática
)""")

    _add_heading(doc, "12.3 Inicialización idempotente", 2)
    doc.add_paragraph(
        "El método init_db() usa CREATE TABLE IF NOT EXISTS, por lo que es seguro "
        "llamarlo en cada cold start. Una bandera de proceso evita repetir el "
        "trabajo dentro de la misma instancia caliente."
    )

    # 13
    _add_heading(doc, "13. Seguridad", 1)
    _add_heading(doc, "13.1 Credenciales", 2)
    _add_bullet_list(doc, [
        "DATABASE_URL contiene usuario + contraseña; nunca se commitea (está en .gitignore).",
        "En Vercel se configura como Environment Variable cifrada en reposo.",
        "Si se compromete, resetear desde Supabase Project Settings → Database → "
        "Reset database password.",
    ])

    _add_heading(doc, "13.2 Row Level Security (RLS)", 2)
    doc.add_paragraph(
        "Supabase habilita RLS por defecto en tablas creadas vía el Table Editor. "
        "Como en este proyecto la app accede con el rol 'postgres' (vía DATABASE_URL), "
        "bypassea RLS — por eso no está habilitada. Si en el futuro se expone la API "
        "REST de PostgREST con el anon key, será obligatorio habilitar RLS y definir "
        "políticas. El esquema schema.sql incluye un bloque comentado con las "
        "instrucciones."
    )

    _add_heading(doc, "13.3 Validación de entrada", 2)
    _add_bullet_list(doc, [
        "El backend usa SQL parametrizado (psycopg/SQLite) — sin riesgo de inyección.",
        "El parser de Excel valida estructura y tipos antes de persistir.",
        "Los CHECK constraints en la BD bloquean valores fuera de dominio.",
    ])

    _add_heading(doc, "13.4 Cifrado", 2)
    _add_bullet_list(doc, [
        "Conexión TCP/SSL al pooler de Supabase (TLS 1.2+).",
        "Datos en reposo cifrados por Supabase (AES-256).",
        "Cookies de sesión Flask firmadas con SECRET_KEY.",
    ])

    # 14
    _add_heading(doc, "14. Backups, mantenimiento y restauración", 1)
    _add_heading(doc, "14.1 Backups automáticos de Supabase", 2)
    _add_bullet_list(doc, [
        "Plan gratuito: 1 backup diario (retención 7 días).",
        "Plan Pro: backups continuos con Point-In-Time Recovery hasta 30 días.",
        "Disponibles en: Supabase Project → Database → Backups.",
    ])

    _add_heading(doc, "14.2 Backup manual (dump SQL)", 2)
    _add_code(doc, """# Exporta el esquema y datos
pg_dump "$DATABASE_URL" > backup_$(date +%Y%m%d).sql

# Restaura
psql "$DATABASE_URL_NUEVA" < backup_20260519.sql""")

    _add_heading(doc, "14.3 Exportación rápida vía la app", 2)
    doc.add_paragraph(
        "Cada estudio puede exportarse a Excel desde la app (endpoint "
        "/api/estudios/<id>/excel). Esto sirve como respaldo individual y como "
        "interoperabilidad con sistemas que no leen Postgres."
    )

    # 15
    _add_heading(doc, "15. Consultas útiles", 1)
    doc.add_paragraph("Algunas consultas SQL que pueden ejecutarse desde el SQL Editor "
                      "de Supabase para análisis ad-hoc:")
    _add_heading(doc, "15.1 Estudios por producto", 2)
    _add_code(doc, """SELECT producto, COUNT(*) AS n_estudios
FROM estudios
GROUP BY producto
ORDER BY n_estudios DESC;""")

    _add_heading(doc, "15.2 Distribución de tipos de gráfico", 2)
    _add_code(doc, """SELECT tipo_grafico, COUNT(*) AS n
FROM estudios
GROUP BY tipo_grafico
ORDER BY n DESC;""")

    _add_heading(doc, "15.3 Estudios con más muestras", 2)
    _add_code(doc, """SELECT e.id, e.nombre, e.producto, COUNT(m.id) AS n_muestras
FROM estudios e
LEFT JOIN muestras m ON m.estudio_id = e.id
GROUP BY e.id
ORDER BY n_muestras DESC
LIMIT 10;""")

    _add_heading(doc, "15.4 Trabajo de un analista en un periodo", 2)
    _add_code(doc, """SELECT id, nombre, producto, fecha_creacion
FROM estudios
WHERE analista = 'Juan Pérez'
  AND fecha_creacion >= '2026-01-01'
  AND fecha_creacion <  '2026-07-01'
ORDER BY fecha_creacion DESC;""")

    _add_heading(doc, "15.5 Lotes con mayor número de subgrupos por producto", 2)
    _add_code(doc, """SELECT e.producto, e.lote, COUNT(m.id) AS n_subgrupos
FROM estudios e
JOIN muestras m ON m.estudio_id = e.id
WHERE e.lote IS NOT NULL
GROUP BY e.producto, e.lote
ORDER BY n_subgrupos DESC
LIMIT 25;""")

    _add_heading(doc, "15.6 Acceso al campo JSONB", 2)
    doc.add_paragraph(
        "Para inspeccionar el contenido del campo 'valores', se pueden usar los "
        "operadores nativos de JSONB:"
    )
    _add_code(doc, """-- Primer valor del subgrupo
SELECT subgrupo, (valores->0)::numeric AS primera_medicion
FROM muestras WHERE estudio_id = 1;

-- Suma de valores en un subgrupo (variables)
SELECT subgrupo,
       (SELECT SUM(v::numeric) FROM jsonb_array_elements_text(valores) AS v) AS suma
FROM muestras WHERE estudio_id = 1;""")

    # 16
    _add_heading(doc, "16. Migración y portabilidad", 1)
    _add_heading(doc, "16.1 De SQLite local a Supabase Postgres", 2)
    _add_numbered_list(doc, [
        "Ejecutar supabase/schema.sql en el SQL Editor.",
        "Configurar DATABASE_URL en el entorno local (.env) o en Vercel.",
        "Migración manual de datos (si aplica): exportar a Excel desde la app, "
        "subirlos en el entorno de producción.",
    ])

    _add_heading(doc, "16.2 Cambiar de proveedor (Neon, Vercel Postgres, RDS, etc.)", 2)
    doc.add_paragraph(
        "El esquema es PostgreSQL estándar; cualquier proveedor con Postgres 13+ "
        "funciona. Solo se necesita actualizar la DATABASE_URL. Si se abandona el "
        "modelo pooler, se puede activar de nuevo prepared_threshold > 0."
    )

    _add_heading(doc, "16.3 Cambiar de RDBMS (MySQL, MariaDB)", 2)
    _add_bullet_list(doc, [
        "El tipo JSONB es específico de Postgres. En MySQL se usaría JSON (similar).",
        "BIGSERIAL pasaría a BIGINT AUTO_INCREMENT.",
        "TIMESTAMPTZ pasaría a DATETIME(6) o TIMESTAMP con manejo manual de zona.",
        "El driver cambiaría de psycopg a mysqlclient o aiomysql.",
    ])

    # 17
    _add_heading(doc, "17. Métricas y monitoreo", 1)
    _add_heading(doc, "17.1 Panel de Supabase", 2)
    _add_bullet_list(doc, [
        ("Database → Database Health: ", "uso de CPU, RAM, conexiones activas."),
        ("Database → Reports: ", "queries más lentos, tablas más grandes, "
         "índices subutilizados."),
        ("Database → Roles: ", "conexiones por rol/usuario."),
        ("Logs → Postgres logs: ", "errores y warnings del motor."),
    ])

    _add_heading(doc, "17.2 Consultas de monitoreo", 2)
    _add_code(doc, """-- Tamaño total de la BD
SELECT pg_size_pretty(pg_database_size(current_database()));

-- Tamaño por tabla (datos + índices)
SELECT relname AS tabla,
       pg_size_pretty(pg_total_relation_size(C.oid)) AS tamano
FROM pg_class C
LEFT JOIN pg_namespace N ON N.oid = C.relnamespace
WHERE nspname = 'public'
  AND C.relkind = 'r'
ORDER BY pg_total_relation_size(C.oid) DESC;

-- Conexiones activas
SELECT count(*) FROM pg_stat_activity
WHERE datname = current_database();""")

    # 18
    _add_heading(doc, "18. Limitaciones conocidas", 1)
    _add_bullet_list(doc, [
        ("Sin autenticación: ", "cualquiera con acceso a la URL pública puede crear, "
         "ver o eliminar estudios. Mitigable con Supabase Auth + RLS."),
        ("Plan gratuito de Supabase: ", "500 MB; suficiente para miles de estudios, "
         "pero conviene monitorear si crece mucho."),
        ("Sin auditoría: ", "no se registra quién creó/modificó cada estudio (más "
         "allá del campo analista, que es de texto libre)."),
        ("Una sola región: ", "el pooler está fijo a una región AWS; usuarios lejanos "
         "experimentarán latencia mayor."),
        ("Conexiones por petición: ", "no se reutilizan dentro del proceso Vercel. "
         "El pooler lo compensa, pero introduce ~10-30 ms por petición."),
    ])

    # 19
    _add_heading(doc, "19. Glosario", 1)
    glosario = [
        ("DDL", "Data Definition Language. Sentencias que definen esquemas (CREATE, ALTER, DROP)."),
        ("DML", "Data Manipulation Language. SELECT, INSERT, UPDATE, DELETE."),
        ("Pooler", "Componente intermedio que gestiona conexiones a la base de datos."),
        ("PgBouncer", "Pooler de conexiones para PostgreSQL usado por Supabase."),
        ("Prepared statement", "Consulta precompilada por el servidor para ejecución repetida."),
        ("RLS", "Row Level Security. Filtra filas según el usuario autenticado."),
        ("JSONB", "JSON binario, indexable y eficiente, en PostgreSQL."),
        ("Transaction Pooler", "Modo del pooler que asigna conexiones por transacción."),
        ("Session Pooler", "Modo del pooler que mantiene la conexión durante toda la sesión."),
        ("CASCADE", "Operación que se propaga a filas relacionadas (ON DELETE CASCADE)."),
        ("FK / PK", "Foreign Key / Primary Key — clave foránea / clave primaria."),
    ]
    for term, desc in glosario:
        p = doc.add_paragraph()
        run = p.add_run(term + ": ")
        run.bold = True; run.font.color.rgb = COLOR_BLUE
        p.add_run(desc)

    out = os.path.join(DOCS, "Informe_Base_Datos.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    print("Generando Informe de Base de Datos…")
    path = generar()
    size_kb = os.path.getsize(path) / 1024
    print(f"  ✓ {os.path.basename(path)} ({size_kb:.1f} KB)")
    print("Listo.")
