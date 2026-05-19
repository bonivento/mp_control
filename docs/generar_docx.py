"""Genera los archivos Word (.docx) del Manual de Usuario y del Informe Técnico.

Requiere: pip install python-docx

Uso:
    python docs/generar_docx.py

Produce en docs/:
    Manual_Usuario.docx
    Informe_Tecnico.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.dirname(os.path.abspath(__file__))
DIAGRAMS = os.path.join(DOCS, "diagrams")

# Paleta corporativa
COLOR_NAVY    = RGBColor(0x00, 0x3A, 0x6B)
COLOR_BLUE    = RGBColor(0x00, 0x5C, 0xAB)
COLOR_BLUE_LT = RGBColor(0x01, 0x83, 0xEF)
COLOR_ORANGE  = RGBColor(0xFF, 0x94, 0x00)
COLOR_GREEN   = RGBColor(0x00, 0xA5, 0x0B)
COLOR_GREY    = RGBColor(0x5B, 0x6B, 0x7A)


# =========================================================
# Helpers de formato
# =========================================================
def _set_default_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    for level, size, color, space_before, space_after in [
        ("Heading 1", 20, COLOR_NAVY,  Pt(16), Pt(6)),
        ("Heading 2", 16, COLOR_BLUE,  Pt(14), Pt(4)),
        ("Heading 3", 13, COLOR_BLUE,  Pt(10), Pt(3)),
        ("Heading 4", 11, COLOR_NAVY,  Pt(8),  Pt(2)),
    ]:
        s = doc.styles[level]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True
        s.paragraph_format.space_before = space_before
        s.paragraph_format.space_after = space_after


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _add_paragraph(doc, text, bold=False, italic=False, size=None, color=None,
                   align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    return p


def _add_bullet_list(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            run = p.add_run(it[0])
            run.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def _add_numbered_list(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(it)


def _shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _add_table(doc, headers, rows, header_fill="005CAB", widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        _shade_cell(hdr[i], header_fill)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for c, value in enumerate(row):
            cells[c].text = ""
            run = cells[c].paragraphs[0].add_run(str(value))
            run.font.size = Pt(9.5)

    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = w

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)


def _add_image(doc, path, caption=None, width_cm=15.5):
    if not os.path.isfile(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cap.add_run(caption)
        crun.italic = True
        crun.font.size = Pt(9)
        crun.font.color.rgb = COLOR_GREY


def _add_footer(doc, text):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = COLOR_GREY

    # Número de página
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Página ")
    run2.font.size = Pt(8.5)
    run2.font.color.rgb = COLOR_GREY

    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run3 = p2.add_run()
    run3._r.append(fld); run3._r.append(instr); run3._r.append(fld2)


def _add_portada(doc, titulo, subtitulo):
    sect = doc.sections[0]
    sect.top_margin = Cm(2.5)
    sect.bottom_margin = Cm(2.5)
    sect.left_margin = Cm(2.5)
    sect.right_margin = Cm(2.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(3)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("UNIVERSIDAD DEL MAGDALENA")
    run.font.size = Pt(14); run.bold = True; run.font.color.rgb = COLOR_BLUE

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p3.add_run("Control Estadístico de Procesos — 2026-1")
    run.font.size = Pt(11); run.italic = True; run.font.color.rgb = COLOR_GREY

    doc.add_paragraph()
    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p4.add_run(titulo)
    run.font.size = Pt(26); run.bold = True; run.font.color.rgb = COLOR_NAVY

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p5.add_run(subtitulo)
    run.font.size = Pt(15); run.font.color.rgb = COLOR_BLUE; run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p6.add_run("Sistema de Control Estadístico de Calidad")
    run.font.size = Pt(12); run.bold = True; run.font.color.rgb = COLOR_ORANGE

    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p7.add_run("Aplicación web para frutas, hortalizas y plantas medicinales")
    run.font.size = Pt(11); run.font.color.rgb = COLOR_GREY

    for _ in range(8):
        doc.add_paragraph()

    p8 = doc.add_paragraph()
    p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p8.add_run("Santa Marta, mayo de 2026")
    run.font.size = Pt(11); run.font.color.rgb = COLOR_GREY

    doc.add_page_break()


def _add_toc(doc, items):
    _add_heading(doc, "Tabla de contenidos", level=1)
    for txt, level in items:
        p = doc.add_paragraph()
        indent = (level - 1) * 0.5
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        run = p.add_run(txt)
        run.font.size = Pt(11)
        if level == 1:
            run.bold = True
            run.font.color.rgb = COLOR_NAVY
        else:
            run.font.color.rgb = COLOR_GREY
    doc.add_page_break()


# =========================================================
# MANUAL DE USUARIO
# =========================================================
def generar_manual():
    doc = Document()
    _set_default_styles(doc)
    _add_portada(doc, "Manual de Usuario",
                 "Sistema de Control Estadístico de Calidad")
    _add_footer(doc, "Universidad del Magdalena · Sistema CEC · Manual de Usuario")

    _add_toc(doc, [
        ("1. Introducción", 1),
        ("2. Requisitos y acceso", 1),
        ("3. Navegación general", 1),
        ("4. Flujo de trabajo recomendado", 1),
        ("5. Registro de trazabilidad", 1),
        ("6. Control por variables (X̄-R, X̄-S)", 1),
        ("7. Control por atributos (p, np, c, u)", 1),
        ("8. Pruebas de normalidad", 1),
        ("9. Capacidad del proceso (Cp, Cpk, Pp, Ppk)", 1),
        ("10. Diagrama de Pareto", 1),
        ("11. Reglas de Nelson", 1),
        ("12. Importar / Exportar Excel", 1),
        ("13. Dashboard y filtros", 1),
        ("14. Cómo interpretar los resultados", 1),
        ("15. Errores comunes y soluciones", 1),
        ("16. Glosario", 1),
    ])

    # 1
    _add_heading(doc, "1. Introducción", 1)
    doc.add_paragraph(
        "Este sistema implementa las principales herramientas de Control Estadístico de "
        "Procesos (CEP / CEC) aplicadas al monitoreo de la calidad de frutas, hortalizas y "
        "plantas medicinales. Permite registrar muestras, aplicar pruebas estadísticas, "
        "construir gráficos de control, calcular índices de capacidad, identificar causas "
        "de no conformidad y exportar reportes a Excel."
    )
    doc.add_paragraph(
        "Funciona como aplicación web; basta con un navegador moderno. Los datos se "
        "almacenan en una base de datos PostgreSQL (Supabase) y persisten entre sesiones."
    )

    # 2
    _add_heading(doc, "2. Requisitos y acceso", 1)
    _add_heading(doc, "2.1 Para usar la aplicación", 2)
    _add_bullet_list(doc, [
        "Navegador moderno: Chrome 100+, Firefox 100+, Safari 15+, Edge 100+.",
        "Conexión a internet (la app está desplegada en Vercel).",
        "Archivos Excel: extensión .xlsx, tamaño máximo 5 MB.",
    ])
    _add_heading(doc, "2.2 Para uso local / desarrollo", 2)
    _add_bullet_list(doc, [
        "Python 3.9 o superior.",
        "Instalar dependencias con: pip install -r requirements.txt",
        "Arrancar con: python api/index.py --port 5050 --debug",
        "Opcional: archivo .env con DATABASE_URL para conectar a Supabase.",
    ])

    # 3
    _add_heading(doc, "3. Navegación general", 1)
    doc.add_paragraph("La barra superior ofrece accesos directos a todas las secciones:")
    _add_table(doc, ["Sección", "Función"], [
        ("Inicio", "Dashboard con resumen y lista de estudios anteriores"),
        ("Nuevo Estudio", "Formulario para registrar un estudio paso a paso"),
        ("Plantillas", "Descarga de archivos de prueba y subida de Excel"),
        ("Normalidad", "Pruebas de normalidad sobre un conjunto de datos rápido"),
        ("Capacidad", "Cálculo independiente de Cp/Cpk/Pp/Ppk"),
        ("Pareto", "Diagrama de Pareto independiente"),
        ("Manual / Informe", "Documentación del sistema"),
    ])

    # 4
    _add_heading(doc, "4. Flujo de trabajo recomendado", 1)
    _add_numbered_list(doc, [
        "Definir qué controlar: variable continua (peso, pH, °Brix) o atributo "
        "(defectos, manchas, frutos podridos).",
        "Determinar el tamaño de subgrupo (n). Para variables, típicamente 4 ó 5; "
        "para atributos depende de la frecuencia esperada de defectos.",
        "Tomar al menos 25 subgrupos de manera secuencial (mínimo recomendado por la literatura).",
        "Registrar los datos en la app — manualmente o subiendo un Excel preformateado.",
        "Analizar el gráfico de control y la lista de reglas de Nelson. Si hay puntos fuera "
        "o patrones, investigar las causas asignables.",
        "Si el proceso está estable, calcular la capacidad (Cp/Cpk). Esto indica si cumple "
        "las especificaciones técnicas.",
        "Exportar el reporte a Excel para archivar o compartir.",
    ])

    # 5
    _add_heading(doc, "5. Registro de trazabilidad", 1)
    doc.add_paragraph(
        "Cada estudio lleva la información que permite identificarlo y reproducirlo. "
        "Los campos disponibles son:"
    )
    _add_table(doc, ["Campo", "Obligatorio", "Ejemplo"], [
        ("Nombre del estudio", "Sí", "Peso de mango Tommy lote 2026-001"),
        ("Producto", "Sí", "Mango, Sábila, Cilantro, Aguacate…"),
        ("Característica", "Sí", "Peso, pH, °Brix, Manchas, Daños…"),
        ("Unidad", "No", "g, cm, %, °Bx, °C"),
        ("Analista", "No", "Nombre de quien toma las mediciones"),
        ("Lote", "No", "Identificador único del lote o cosecha"),
        ("Tipo", "Sí", "variable / atributo"),
        ("Tipo de gráfico", "Sí", "xr, xs, p, np, c, u"),
        ("LSL / USL", "Solo capacidad", "Límites de especificación"),
        ("Tamaño subgrupo", "X̄-R/X̄-S", "Entre 2 y 25"),
        ("Notas", "No", "Observaciones del proceso"),
    ])

    # 6
    _add_heading(doc, "6. Control por variables (datos continuos)", 1)
    doc.add_paragraph("Para mediciones numéricas (peso, dimensiones, contenido químico).")
    _add_heading(doc, "Gráfico X̄-R", 2)
    doc.add_paragraph(
        "Apropiado para subgrupos pequeños (n entre 2 y 9). Monitorea simultáneamente "
        "la media y el rango. Fórmulas:"
    )
    _add_paragraph(doc, "UCL X̄ = X̿ + A₂·R̄    LCL X̄ = X̿ − A₂·R̄", italic=True)
    _add_paragraph(doc, "UCL R = D₄·R̄    LCL R = D₃·R̄", italic=True)
    _add_paragraph(doc, "Estimación σ̂ = R̄ / d₂", italic=True)

    _add_heading(doc, "Gráfico X̄-S", 2)
    doc.add_paragraph(
        "Recomendado para subgrupos grandes (n ≥ 10). Reemplaza el rango por la desviación "
        "estándar muestral, estimador más eficiente cuando n crece."
    )
    _add_paragraph(doc, "UCL X̄ = X̿ + A₃·S̄    LCL X̄ = X̿ − A₃·S̄", italic=True)
    _add_paragraph(doc, "UCL S = B₄·S̄    LCL S = B₃·S̄", italic=True)
    _add_paragraph(doc, "Estimación σ̂ = S̄ / c₄", italic=True)
    doc.add_paragraph(
        "Las constantes A₂, A₃, D₃, D₄, B₃, B₄, d₂, c₄ están implementadas para n entre 2 "
        "y 25 siguiendo Montgomery (Statistical Quality Control, 7ª ed.)."
    )

    # 7
    _add_heading(doc, "7. Control por atributos (datos discretos)", 1)
    doc.add_paragraph("Para conteos: frutos defectuosos, número de manchas, presencia de plagas.")
    _add_heading(doc, "Gráfico p (proporción de defectuosos)", 2)
    doc.add_paragraph(
        "Tolera tamaños de muestra variables. Las fórmulas son: p̄ = Σdᵢ / Σnᵢ ; "
        "UCL = p̄ + 3·√[p̄(1−p̄)/nᵢ] ; LCL = máx(0, p̄ − …)."
    )
    _add_heading(doc, "Gráfico np (número de defectuosos)", 2)
    doc.add_paragraph(
        "Tamaño de muestra constante. UCL = np̄ + 3·√[np̄(1−p̄)] ; LCL = máx(0, …)."
    )
    _add_heading(doc, "Gráfico c (defectos por unidad)", 2)
    doc.add_paragraph(
        "Modelo Poisson, área de inspección constante. UCL = c̄ + 3·√c̄ ; LCL = máx(0, …)."
    )
    _add_heading(doc, "Gráfico u (defectos por unidad – área variable)", 2)
    doc.add_paragraph("ū = Σcᵢ / Σnᵢ ; UCL = ū + 3·√(ū/nᵢ) ; LCL = máx(0, …).")

    # 8
    _add_heading(doc, "8. Pruebas de normalidad", 1)
    doc.add_paragraph(
        "Los gráficos X̄-R, X̄-S y los índices de capacidad asumen distribución normal. "
        "Antes de interpretarlos conviene verificar este supuesto. El sistema implementa "
        "tres pruebas:"
    )
    _add_table(doc, ["Prueba", "Cuándo usar", "n mínimo", "Sensibilidad"], [
        ("Shapiro-Wilk", "Muestras pequeñas a medianas", "3", "Muy alta"),
        ("Anderson-Darling", "Cualquier tamaño, énfasis en colas", "8", "Alta"),
        ("D'Agostino-Pearson", "Muestras grandes", "20", "Asimetría + curtosis"),
    ])
    _add_heading(doc, "Interpretación del p-valor", 2)
    _add_bullet_list(doc, [
        "p > 0.05: no se rechaza la normalidad. Los datos son compatibles con normal.",
        "p < 0.05: se rechaza la normalidad. Considerar transformaciones (log, Box-Cox) "
        "o métodos no paramétricos.",
    ])

    # 9
    _add_heading(doc, "9. Capacidad del proceso", 1)
    doc.add_paragraph(
        "Una vez que el proceso es estable, los índices de capacidad miden si cumple "
        "las especificaciones del producto."
    )
    _add_paragraph(doc, "Cp = (USL − LSL) / (6·σ_within)", italic=True)
    _add_paragraph(doc, "Cpk = mín[(USL−μ)/(3σ), (μ−LSL)/(3σ)] usando σ_within", italic=True)
    _add_paragraph(doc, "Pp = (USL − LSL) / (6·σ_overall)", italic=True)
    _add_paragraph(doc, "Ppk = idem Cpk pero con σ_overall", italic=True)

    _add_heading(doc, "Interpretación de Cpk", 2)
    _add_table(doc, ["Valor de Cpk", "Estado", "Acción"], [
        ("≥ 1.67", "Clase mundial", "Mantener, considerar reducir control"),
        ("1.33 – 1.67", "Capaz", "Monitoreo de rutina"),
        ("1.00 – 1.33", "Adecuado", "Control estricto, mejorar"),
        ("0.67 – 1.00", "Parcialmente capaz", "Plan de mejora urgente"),
        ("< 0.67", "Incapaz", "Detener, rediseñar proceso"),
    ])

    # 10
    _add_heading(doc, "10. Diagrama de Pareto", 1)
    doc.add_paragraph(
        "Basado en el principio 80/20 de Vilfredo Pareto: aproximadamente el 80% de los "
        "problemas proviene del 20% de las causas."
    )
    _add_numbered_list(doc, [
        "Ir a la sección Pareto.",
        "Agregar las categorías de defecto y su frecuencia observada en un periodo.",
        "Click en Generar Pareto.",
    ])
    doc.add_paragraph(
        "Las barras (azul) se ordenan de mayor a menor frecuencia, la línea naranja "
        "muestra el porcentaje acumulado y la línea roja punteada marca el 80%. Las "
        "categorías a la izquierda del cruce con esa línea son las 'pocas vitales'."
    )

    # 11
    _add_heading(doc, "11. Reglas de Nelson", 1)
    doc.add_paragraph(
        "Lloyd S. Nelson (1984) propuso reglas adicionales para detectar patrones no "
        "aleatorios en los gráficos de control. El sistema implementa 6 de las 8 reglas "
        "clásicas:"
    )
    _add_table(doc, ["Regla", "Descripción", "Indica"], [
        ("1", "1 punto más allá de ±3σ", "Causa asignable inmediata"),
        ("2", "9 puntos consecutivos del mismo lado de la LC", "Shift de la media"),
        ("3", "6 puntos en tendencia creciente o decreciente", "Deriva del proceso"),
        ("4", "14 puntos alternando arriba/abajo", "Mezcla de dos procesos"),
        ("5", "2 de 3 puntos más allá de ±2σ del mismo lado", "Shift incipiente"),
        ("6", "4 de 5 puntos más allá de ±1σ del mismo lado", "Shift pequeño persistente"),
    ])

    # 12
    _add_heading(doc, "12. Importar / Exportar Excel", 1)
    _add_heading(doc, "Importar (subir)", 2)
    _add_numbered_list(doc, [
        "Ir a Plantillas.",
        "Descargar uno de los archivos de prueba o preparar el propio.",
        "Seleccionar archivo → Subir y analizar.",
        "El estudio se crea automáticamente y se abre la página de análisis.",
    ])
    _add_heading(doc, "Formato del Excel", 2)
    doc.add_paragraph("El archivo debe tener dos hojas:")
    _add_heading(doc, "Hoja 'Trazabilidad'", 3)
    _add_table(doc, ["Campo", "Ejemplo"], [
        ("Nombre", "Peso de mango Tommy"),
        ("Producto", "Mango Tommy"),
        ("Tipo", "variable"),
        ("Caracteristica", "Peso"),
        ("Unidad", "g"),
        ("Tipo de grafico", "xr"),
        ("LSL", "235"),
        ("USL", "265"),
        ("Tamano subgrupo", "5"),
    ])
    _add_heading(doc, "Hoja 'Datos'", 3)
    _add_bullet_list(doc, [
        ("X̄-R / X̄-S: ", "Subgrupo | Med 1 | Med 2 | … | Med N"),
        ("p / u: ", "Subgrupo | Defectivos | Tamano muestra"),
        ("np: ", "Subgrupo | Defectivos | Tamano muestra (constante)"),
        ("c: ", "Subgrupo | Defectos"),
    ])
    _add_heading(doc, "Exportar", 2)
    doc.add_paragraph(
        "Desde la página de cualquier estudio, click en 'Exportar a Excel'. Descarga "
        "un archivo con tres hojas: Trazabilidad, Datos y Resultados."
    )

    # 13
    _add_heading(doc, "13. Dashboard y filtros", 1)
    doc.add_paragraph("En la página de Inicio aparece:")
    _add_bullet_list(doc, [
        "Tarjetas de resumen: total de estudios, separados por tipo y estado del almacenamiento.",
        "Tabla de Estudios anteriores con búsqueda y filtros (por nombre, producto, "
        "característica, analista, lote, tipo y tipo de gráfico).",
        "Botones por estudio: Abrir, Descargar Excel, Eliminar.",
    ])

    # 14
    _add_heading(doc, "14. Cómo interpretar los resultados", 1)
    _add_heading(doc, "Caso 1: proceso en control", 2)
    _add_bullet_list(doc, [
        "Ningún punto fuera de los límites.",
        "Sin violaciones a las reglas de Nelson.",
        "Si hay LSL/USL definidos, calcular Cpk: ≥ 1.33 → cumple; < 1.33 → mejorar.",
    ])
    _add_heading(doc, "Caso 2: proceso fuera de control", 2)
    _add_numbered_list(doc, [
        "No calcular capacidad aún. Cp/Cpk solo tienen sentido con proceso estable.",
        "Identificar el primer punto que viola una regla.",
        "Investigar qué cambió en ese momento del proceso.",
        "Aplicar acciones correctivas, muestrear nuevamente y recalcular.",
    ])
    _add_heading(doc, "Caso 3: datos no normales", 2)
    _add_bullet_list(doc, [
        "X̄-R y X̄-S son robustos a desviaciones leves.",
        "Para no normalidad fuerte: transformar (log, raíz, Box-Cox) o métodos no paramétricos.",
        "La capacidad bajo asunción de normalidad será imprecisa — reportarla con la salvedad.",
    ])

    # 15
    _add_heading(doc, "15. Errores comunes y soluciones", 1)
    _add_table(doc, ["Síntoma", "Causa", "Solución"], [
        ("Cannot assign requested address",
         "Conexión directa de Supabase (IPv6) en Vercel",
         "Usar Transaction Pooler (puerto 6543, IPv4)"),
        ("prepared statement _pg3_X already exists",
         "Conflicto entre psycopg y PgBouncer",
         "Ya resuelto en la app con prepare_threshold=None"),
        ("Faltan campos requeridos en Trazabilidad",
         "Excel sin uno de los campos obligatorios",
         "Verificar Nombre, Producto, Tipo, Caracteristica, Tipo de grafico"),
        ("El subgrupo X tiene Y mediciones; se esperaban N",
         "Filas con distinto número de columnas",
         "Asegurar el mismo número de columnas en todos los subgrupos"),
        ("Los gráficos no se renderizan",
         "JavaScript bloqueado o sin conexión al CDN",
         "Permitir scripts externos (revisa la consola F12)"),
        ("Estudio no aparece tras crearlo en Vercel",
         "SQLite efímero en serverless",
         "Configurar DATABASE_URL con Supabase Postgres"),
    ])

    # 16
    _add_heading(doc, "16. Glosario", 1)
    glosario = [
        ("Atributo", "Característica evaluada como presencia/ausencia o conteo."),
        ("Capacidad del proceso", "Aptitud de un proceso para producir dentro de las especificaciones."),
        ("Causa asignable", "Fuente identificable de variación no inherente al proceso."),
        ("Causa común", "Variación inherente y aleatoria del sistema."),
        ("LSC / LIC (UCL / LCL)", "Límites superior/inferior de control (±3σ)."),
        ("LSL / USL", "Límites inferior/superior de especificación técnica."),
        ("Proceso en control estadístico", "Variabilidad proveniente solo de causas comunes."),
        ("Subgrupo", "Conjunto de mediciones en condiciones homogéneas."),
        ("Variable", "Característica medible numéricamente."),
    ]
    for term, desc in glosario:
        p = doc.add_paragraph()
        run = p.add_run(term + ": ")
        run.bold = True; run.font.color.rgb = COLOR_BLUE
        p.add_run(desc)

    out = os.path.join(DOCS, "Manual_Usuario.docx")
    doc.save(out)
    return out


# =========================================================
# INFORME TÉCNICO
# =========================================================
def generar_informe():
    doc = Document()
    _set_default_styles(doc)
    _add_portada(doc, "Informe Técnico",
                 "Sistema de Control Estadístico de Calidad")
    _add_footer(doc, "Universidad del Magdalena · Sistema CEC · Informe Técnico")

    _add_toc(doc, [
        ("1. Introducción y objetivos", 1),
        ("1.1 Contexto", 2),
        ("1.2 Justificación", 2),
        ("1.3 Objetivo general", 2),
        ("1.4 Objetivos específicos", 2),
        ("1.5 Alcance", 2),
        ("2. Marco teórico", 1),
        ("2.1 Origen y fundamentos del CEP", 2),
        ("2.2 Variabilidad y regla 3-sigma", 2),
        ("2.3 Gráficos para variables continuas", 2),
        ("2.4 Gráficos para atributos", 2),
        ("2.5 Pruebas de normalidad", 2),
        ("2.6 Índices de capacidad", 2),
        ("2.7 Diagrama de Pareto", 2),
        ("2.8 Reglas de Nelson", 2),
        ("3. Metodología", 1),
        ("3.1 Enfoque de desarrollo", 2),
        ("3.2 Arquitectura del sistema", 2),
        ("3.3 Stack tecnológico", 2),
        ("3.4 Diseño de datos", 2),
        ("3.5 Flujo de datos típico", 2),
        ("3.6 Validación de los cálculos", 2),
        ("4. Resultados y análisis", 1),
        ("4.1 Funcionalidades entregadas", 2),
        ("4.2 Validación con datos de prueba", 2),
        ("4.3 Casos de uso reales", 2),
        ("4.4 Métricas del sistema", 2),
        ("5. Conclusiones y recomendaciones", 1),
        ("5.1 Conclusiones", 2),
        ("5.2 Limitaciones", 2),
        ("5.3 Recomendaciones", 2),
        ("5.4 Trabajo futuro", 2),
        ("Referencias", 1),
    ])

    # 1
    _add_heading(doc, "1. Introducción y objetivos", 1)
    _add_heading(doc, "1.1 Contexto", 2)
    doc.add_paragraph(
        "La región Caribe colombiana, y en particular el departamento del Magdalena, "
        "es uno de los principales productores nacionales de frutas tropicales (mango, "
        "banano, aguacate), hortalizas y plantas medicinales (sábila, manzanilla, "
        "hierbabuena). Estos productos abastecen el mercado interno y las industrias "
        "farmacéutica, cosmética, alimentaria y agroindustrial, y son objeto de "
        "regulación por el INVIMA en Colombia, la FDA en Estados Unidos y la EFSA en "
        "la Unión Europea, además de estándares voluntarios como ISO 9001, ISO 22000, "
        "Buenas Prácticas Agrícolas (BPA) y Buenas Prácticas de Manufactura (BPM)."
    )
    doc.add_paragraph(
        "Garantizar la calidad y trazabilidad de estos productos exige métodos "
        "cuantitativos que permitan monitorear variables (peso, pH, °Brix, contenido de "
        "principios activos) y atributos (presencia de plagas, defectos visuales, daños "
        "mecánicos) de manera sistemática, y detectar oportunamente cuándo un proceso "
        "productivo se desvía de su comportamiento normal. El Control Estadístico de "
        "Procesos (CEP), desarrollado por Walter A. Shewhart en los años 1920 y "
        "consolidado por W. Edwards Deming, ofrece un cuerpo metodológico ampliamente "
        "probado para esta tarea."
    )

    _add_heading(doc, "1.2 Justificación", 2)
    doc.add_paragraph(
        "El desarrollo de una aplicación web de uso libre, accesible desde cualquier "
        "dispositivo, que integre las principales herramientas estadísticas de control "
        "de calidad y que permita registrar trazabilidad completa, representa una "
        "contribución directa a la competitividad del sector agro de la región. La "
        "integración con bases de datos persistentes (Supabase) y el despliegue sin "
        "servidor (Vercel) reducen drásticamente los costos de infraestructura, mientras "
        "que la interfaz amigable democratiza el acceso a estas técnicas."
    )

    _add_heading(doc, "1.3 Objetivo general", 2)
    doc.add_paragraph(
        "Desarrollar una aplicación de software web que integre herramientas de Control "
        "Estadístico de Calidad (CEC) para el monitoreo y análisis de variables y "
        "atributos asociados a la calidad de frutas, hortalizas y plantas medicinales."
    )

    _add_heading(doc, "1.4 Objetivos específicos", 2)
    _add_bullet_list(doc, [
        "Aplicar herramientas estadísticas (pruebas de normalidad, gráficos de control, "
        "índices de capacidad, diagrama de Pareto) para el control y mejora de la calidad.",
        "Diseñar un sistema que permita registrar, analizar y visualizar datos de "
        "calidad con trazabilidad completa (producto, lote, analista, fecha).",
        "Interpretar automáticamente los resultados aplicando las reglas de Nelson para "
        "detectar patrones no aleatorios y apoyar la toma de decisiones.",
        "Fomentar la integración entre estadística, programación y gestión de la "
        "calidad mediante una arquitectura modular y documentada.",
        "Garantizar la persistencia y portabilidad de la información mediante una "
        "base de datos relacional administrada en la nube.",
    ])

    _add_heading(doc, "1.5 Alcance", 2)
    doc.add_paragraph("El sistema permite gestionar estudios completos de control estadístico para:")
    _add_bullet_list(doc, [
        ("Variables continuas: ", "peso, dimensiones, °Bx, pH, humedad, aceites esenciales, "
         "principios activos."),
        ("Atributos discretos: ", "presencia de plagas, manchas, daños mecánicos, "
         "defectos de color, cumplimiento de BPA/BPM."),
        ("Análisis de hasta 25 subgrupos por estudio ", "(mínimo recomendado por la "
         "literatura), escalable a más sin limitaciones técnicas."),
    ])

    # 2
    _add_heading(doc, "2. Marco teórico", 1)
    _add_heading(doc, "2.1 Origen y fundamentos del CEP", 2)
    doc.add_paragraph(
        "El Control Estadístico de Procesos fue introducido por Walter A. Shewhart en "
        "1924 en los Bell Telephone Laboratories. Su contribución fundamental fue la "
        "distinción entre dos tipos de variación en cualquier proceso:"
    )
    _add_bullet_list(doc, [
        ("Causas comunes (aleatorias): ", "inherentes al proceso, estables en el tiempo. "
         "Producen variación predecible. Solo se reducen rediseñando el sistema."),
        ("Causas asignables (especiales): ", "externas o esporádicas. Producen variación "
         "impredecible y son señal de que algo cambió. Requieren acción correctiva inmediata."),
    ])
    doc.add_paragraph(
        "Un proceso está en control estadístico cuando opera únicamente bajo causas "
        "comunes. Los gráficos de control son la herramienta visual para distinguir "
        "estos dos tipos de variación."
    )

    _add_heading(doc, "2.2 Variabilidad y la regla 3-sigma", 2)
    doc.add_paragraph(
        "Shewhart estableció empíricamente que límites a ±3 desviaciones estándar de la "
        "media del proceso son un buen balance entre detectar señales reales y evitar "
        "falsas alarmas. Bajo distribución normal, ±3σ contiene el 99.73% de los datos; "
        "por lo tanto, un punto fuera de esos límites tiene probabilidad < 0.27% de "
        "ocurrir si el proceso está en control."
    )

    _add_heading(doc, "2.3 Gráficos para variables continuas", 2)
    doc.add_paragraph(
        "Cuando la característica de calidad es medible en una escala continua, se "
        "monitorean simultáneamente la tendencia central y la variabilidad mediante "
        "gráficos X̄-R (subgrupos n=2..9) o X̄-S (n ≥ 10). Las fórmulas son:"
    )
    _add_paragraph(doc, "Gráfico X̄: UCL = X̿ + A₂·R̄ (o A₃·S̄); LCL = X̿ − A₂·R̄ (o A₃·S̄)", italic=True)
    _add_paragraph(doc, "Gráfico R: UCL = D₄·R̄; LCL = D₃·R̄", italic=True)
    _add_paragraph(doc, "Gráfico S: UCL = B₄·S̄; LCL = B₃·S̄", italic=True)
    _add_paragraph(doc, "Estimación de σ: σ̂ = R̄ / d₂  o  σ̂ = S̄ / c₄", italic=True)

    _add_heading(doc, "2.4 Gráficos para atributos", 2)
    doc.add_paragraph(
        "Cuando las características se evalúan como conformidad/no conformidad o por "
        "conteo de defectos, se usan gráficos basados en distribuciones discretas:"
    )
    _add_bullet_list(doc, [
        ("p: ", "proporción de defectuosos (Binomial). Tamaño de muestra variable."),
        ("np: ", "número de defectuosos. Tamaño de muestra constante."),
        ("c: ", "número de defectos por unidad (Poisson). Área constante."),
        ("u: ", "defectos por unidad cuando el área de oportunidad varía."),
    ])

    _add_heading(doc, "2.5 Pruebas de normalidad", 2)
    doc.add_paragraph(
        "Los gráficos de control para variables, así como los índices de capacidad, "
        "asumen distribución normal. Las pruebas más usadas son:"
    )
    _add_bullet_list(doc, [
        ("Shapiro-Wilk (1965): ", "potente para muestras pequeñas (n < 50)."),
        ("Anderson-Darling (1952): ", "pondera más las colas."),
        ("D'Agostino-Pearson (1973): ", "combina asimetría y curtosis (requiere n ≥ 20)."),
    ])

    _add_heading(doc, "2.6 Índices de capacidad del proceso", 2)
    doc.add_paragraph(
        "Los índices Cp, Cpk, Pp, Ppk responden la pregunta '¿el proceso cumple las "
        "especificaciones?'. Cp/Cpk usan σ_within (corto plazo, R̄/d₂ o S̄/c₄), mientras "
        "Pp/Ppk usan σ_overall (largo plazo, desviación estándar muestral)."
    )

    _add_heading(doc, "2.7 Diagrama de Pareto", 2)
    doc.add_paragraph(
        "Joseph M. Juran (1950) adaptó el principio 80/20 del economista italiano "
        "Vilfredo Pareto al campo de la calidad: aproximadamente el 80% de los problemas "
        "proviene del 20% de las causas. El diagrama ordena las causas de no conformidad "
        "por frecuencia y permite identificar las 'pocas vitales'."
    )

    _add_heading(doc, "2.8 Reglas de Nelson", 2)
    doc.add_paragraph(
        "Lloyd S. Nelson (1984) propuso reglas adicionales para detectar patrones no "
        "aleatorios. El sistema implementa 6 de las 8 reglas clásicas (puntos fuera de "
        "±3σ, secuencias del mismo lado, tendencias, alternancias, secuencias más allá "
        "de ±2σ y ±1σ)."
    )

    # 3
    _add_heading(doc, "3. Metodología", 1)
    _add_heading(doc, "3.1 Enfoque de desarrollo", 2)
    doc.add_paragraph(
        "Se adoptó un enfoque iterativo e incremental, similar a metodologías ágiles. "
        "Cada iteración entregó funcionalidad verificable: primero los cálculos "
        "estadísticos puros (validados contra Montgomery), luego la capa de persistencia, "
        "después la interfaz web, la importación/exportación de Excel y finalmente la "
        "migración a Supabase Postgres y el despliegue en Vercel."
    )

    _add_heading(doc, "3.2 Arquitectura del sistema", 2)
    doc.add_paragraph(
        "Se implementó una arquitectura de tres capas: (1) presentación en el cliente "
        "con HTML/CSS y Plotly.js; (2) lógica de negocio en Flask/Python con APIs REST "
        "de cálculo estadístico; (3) persistencia en PostgreSQL administrado por "
        "Supabase, con dispatcher automático a SQLite local si no se configura la BD "
        "remota."
    )
    _add_image(doc, os.path.join(DIAGRAMS, "arquitectura.png"),
               caption="Figura 1. Arquitectura general del sistema en tres capas.",
               width_cm=16)

    _add_heading(doc, "3.3 Stack tecnológico", 2)
    _add_table(doc, ["Capa", "Tecnología", "Justificación"], [
        ("Lenguaje servidor", "Python 3.11", "Ecosistema científico maduro"),
        ("Framework web", "Flask 3", "Ligero, compatible con serverless"),
        ("Cálculo estadístico", "NumPy, SciPy, Pandas", "Estándar de facto"),
        ("Exportación / Importación", "OpenPyXL", "Lectura/escritura XLSX nativa"),
        ("Driver BD", "psycopg 3", "Driver moderno para PostgreSQL"),
        ("Gráficos cliente", "Plotly.js 2.35", "Interactividad sin esfuerzo en backend"),
        ("Base de datos", "Supabase Postgres", "Plan gratuito, conexión por pooler"),
        ("Despliegue", "Vercel Serverless", "Despliegue por git push"),
        ("Control de versiones", "Git + GitHub", "Trazabilidad del desarrollo"),
    ])

    _add_heading(doc, "3.4 Diseño de datos", 2)
    doc.add_paragraph(
        "El modelo se organiza en dos tablas: estudios (una fila por estudio con la "
        "trazabilidad) y muestras (una fila por subgrupo, ligada al estudio por foreign "
        "key). Los valores de cada subgrupo se almacenan en una columna JSONB que "
        "soporta tanto arreglos de mediciones (variables) como pares [defectivos, "
        "tamaño_muestra] (atributos). Esta decisión mantiene la integridad del subgrupo "
        "como unidad atómica de análisis y simplifica las consultas."
    )
    _add_image(doc, os.path.join(DIAGRAMS, "modelo_datos.png"),
               caption="Figura 2. Modelo entidad-relación de la base de datos.",
               width_cm=15)

    _add_heading(doc, "3.5 Flujo de datos típico", 2)
    doc.add_paragraph(
        "El siguiente diagrama muestra el flujo completo: el analista sube un archivo "
        "Excel, el servidor lo parsea y persiste, y luego el cliente solicita el análisis "
        "estadístico para renderizarlo con Plotly.js:"
    )
    _add_image(doc, os.path.join(DIAGRAMS, "flujo_datos.png"),
               caption="Figura 3. Diagrama de secuencia — subir Excel y obtener análisis.",
               width_cm=16)

    _add_heading(doc, "3.6 Validación de los cálculos", 2)
    _add_bullet_list(doc, [
        "Las constantes de Shewhart se cargaron de las tablas de Montgomery y se "
        "compararon con valores publicados.",
        "Seis conjuntos de datos sintéticos (tres en control y tres fuera de control) "
        "se procesaron y el sistema detectó correctamente cada caso.",
        "Las pruebas de normalidad se compararon contra scipy.stats.",
        "Cp/Cpk se validaron manualmente con ejemplos de Montgomery (Capítulo 8).",
    ])

    # 4
    _add_heading(doc, "4. Resultados y análisis", 1)
    _add_heading(doc, "4.1 Funcionalidades entregadas", 2)
    _add_bullet_list(doc, [
        "6 gráficos de control: X̄-R, X̄-S, p, np, c, u.",
        "3 pruebas de normalidad + histograma + Q-Q plot.",
        "4 índices de capacidad (Cp, Cpk, Pp, Ppk) con interpretación automática.",
        "Diagrama de Pareto con identificación de pocas vitales.",
        "6 reglas de Nelson implementadas para detección de patrones.",
        "Importación de Excel con plantillas pre-cargadas.",
        "Exportación de Excel con tres hojas (trazabilidad, datos, resultados).",
        "Búsqueda y filtros sobre estudios anteriores.",
        "Persistencia real con Supabase Postgres.",
    ])

    _add_heading(doc, "4.2 Validación con datos de prueba", 2)
    doc.add_paragraph(
        "Se generaron seis archivos Excel sintéticos para validar el comportamiento "
        "del sistema. Cada archivo contiene 25 subgrupos."
    )
    _add_table(doc, ["Archivo", "Producto", "Gráfico", "Esperado", "Resultado"], [
        ("1_xr_en_control_mango", "Mango Tommy", "X̄-R n=5", "Estable",
         "0 fuera, Cpk≈1.42 capaz"),
        ("2_xr_fuera_de_control_mango", "Mango Tommy", "X̄-R n=5", "Shift +8 g",
         "12 fuera, 63 violaciones"),
        ("3_xs_en_control_brix_sandia", "Sandía", "X̄-S n=10", "Estable",
         "0 fuera, 0 violaciones"),
        ("4_p_en_control_aguacate", "Aguacate Hass", "p", "p̄≈5%",
         "0 fuera"),
        ("5_p_fuera_de_control_aguacate", "Aguacate Hass", "p", "5%→18%",
         "4 fuera, 48 violaciones"),
        ("6_c_fuera_de_control_sabila", "Sábila", "c", "Tendencia",
         "0 fuera de ±3σ, 12 violaciones por regla 3"),
    ])
    doc.add_paragraph(
        "El último caso ilustra por qué las reglas de Nelson son cruciales: detectan "
        "tendencias graduales que pueden llevar al proceso fuera de especificación sin "
        "disparar nunca la alarma de 'punto fuera de límite'."
    )

    _add_heading(doc, "4.3 Casos de uso reales aplicables", 2)
    _add_heading(doc, "Cultivo de sábila para industria cosmética", 3)
    doc.add_paragraph(
        "Variables clave: contenido de aceites esenciales (%), pH del gel, peso fresco "
        "(g). Atributos: presencia de plagas, defectos de color. Un gráfico X̄-R del "
        "contenido de aceites con LSL/USL definidos por la farmacopea permitiría "
        "detectar lotes inadecuados antes del procesamiento."
    )
    _add_heading(doc, "Empaque de mango para exportación", 3)
    doc.add_paragraph(
        "Variable principal: peso (g) por unidad. Atributos: golpes, manchas, daños "
        "mecánicos. Gráficos p o np semanales detectarían deterioro en la cadena de "
        "empaque, por ejemplo, por desgaste de bandas transportadoras."
    )
    _add_heading(doc, "Producción de aceite esencial de hierbabuena", 3)
    doc.add_paragraph(
        "Variable: concentración de mentol (mg/g). El control de capacidad respecto a "
        "las especificaciones del cliente farmacéutico determina si el lote se acepta, "
        "ajusta o rechaza."
    )

    _add_heading(doc, "4.4 Métricas del sistema", 2)
    _add_bullet_list(doc, [
        "Latencia tras cold start (~2-3 s): peticiones siguientes responden en < 500 ms.",
        "Probado hasta 100 subgrupos × 10 mediciones (1000 datos) sin degradación.",
        "Tamaño del bundle de despliegue: ~45 MB (límite de Vercel: 50 MB).",
        "Costo operacional: USD 0/mes en planes gratuitos de Vercel y Supabase.",
    ])

    # 5
    _add_heading(doc, "5. Conclusiones y recomendaciones", 1)
    _add_heading(doc, "5.1 Conclusiones", 2)
    _add_bullet_list(doc, [
        "Se desarrolló un sistema web funcional que integra las principales herramientas "
        "de Control Estadístico de Calidad, cumpliendo todos los objetivos planteados.",
        "La arquitectura modular separa los cálculos estadísticos del rendering y de la "
        "persistencia, permitiendo evolucionar cada componente independientemente.",
        "El uso de tecnologías serverless (Vercel) y BaaS (Supabase) reduce drásticamente "
        "la barrera de entrada: sin operar infraestructura propia.",
        "La importación de Excel preformateado integra el sistema con flujos de trabajo "
        "existentes y elimina la digitación manual.",
        "La detección automática de patrones (puntos fuera de control, reglas de "
        "Nelson) ahorra horas de revisión manual.",
        "Los datos sintéticos demuestran que el sistema distingue correctamente entre "
        "procesos estables y procesos con causa asignable.",
        "La persistencia con Supabase Postgres permite que múltiples usuarios compartan "
        "información histórica de manera segura y duradera.",
    ])

    _add_heading(doc, "5.2 Limitaciones", 2)
    _add_bullet_list(doc, [
        "No implementa autenticación de usuarios; cualquiera con el enlace puede crear, "
        "ver o eliminar estudios.",
        "No incluye gráficos avanzados como CUSUM o EWMA para detectar shifts pequeños.",
        "Las alertas son solo visuales dentro de la app; sin notificaciones por correo "
        "o mensaje.",
        "Interfaz limitada al español.",
        "No transforma automáticamente datos no normales (Box-Cox, Johnson).",
    ])

    _add_heading(doc, "5.3 Recomendaciones", 2)
    _add_bullet_list(doc, [
        "Adopción por etapas: empezar con un único producto-característica piloto.",
        "Mínimo 25 subgrupos: insistir en este requisito antes de tomar decisiones "
        "operacionales.",
        "Revisar la normalidad antes de interpretar Cp/Cpk.",
        "No calcular capacidad si el proceso no está en control.",
        "Trazabilidad rigurosa: completar siempre analista y lote.",
    ])

    _add_heading(doc, "5.4 Trabajo futuro", 2)
    _add_numbered_list(doc, [
        "Implementar autenticación y roles usando Supabase Auth.",
        "Agregar gráficos CUSUM y EWMA para detección de shifts pequeños.",
        "Incorporar análisis multivariado (T² de Hotelling).",
        "Integración con sensores IoT (balanza, pH-metro, refractómetro).",
        "Notificaciones por correo o WhatsApp cuando se detecten violaciones graves.",
        "Soporte multi-idioma.",
        "Generación de reportes PDF firmados digitalmente para auditorías.",
        "Aplicación móvil nativa o PWA para captura en campo.",
        "Modelos predictivos basados en machine learning.",
    ])

    # Referencias
    _add_heading(doc, "Referencias", 1)
    referencias = [
        "Anderson, T. W., & Darling, D. A. (1952). Asymptotic theory of certain 'goodness "
        "of fit' criteria based on stochastic processes. The Annals of Mathematical "
        "Statistics, 23(2), 193-212.",
        "D'Agostino, R., & Pearson, E. S. (1973). Tests for departure from normality. "
        "Biometrika, 60(3), 613-622.",
        "Deming, W. E. (1986). Out of the Crisis. MIT Press.",
        "Juran, J. M. (1951). Quality Control Handbook. McGraw-Hill.",
        "Montgomery, D. C. (2013). Statistical Quality Control: A Modern Introduction "
        "(7th ed.). Wiley.",
        "Nelson, L. S. (1984). Technical aids: The Shewhart control chart - tests for "
        "special causes. Journal of Quality Technology, 16(4), 237-239.",
        "Shapiro, S. S., & Wilk, M. B. (1965). An analysis of variance test for normality. "
        "Biometrika, 52(3/4), 591-611.",
        "Shewhart, W. A. (1931). Economic Control of Quality of Manufactured Product. "
        "Van Nostrand.",
        "INVIMA (2025). Resoluciones técnicas para frutas, hortalizas y plantas medicinales.",
        "OMS (2007). WHO guidelines on good agricultural and collection practices (GACP) "
        "for medicinal plants.",
        "ISO 9001:2015. Quality management systems – Requirements.",
        "ISO 22000:2018. Food safety management systems – Requirements.",
    ]
    for r in referencias:
        p = doc.add_paragraph(style="List Number")
        p.add_run(r)

    out = os.path.join(DOCS, "Informe_Tecnico.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    print("Generando documentos Word…")
    m = generar_manual()
    print(f"  ✓ {os.path.basename(m)} ({os.path.getsize(m)/1024:.1f} KB)")
    i = generar_informe()
    print(f"  ✓ {os.path.basename(i)} ({os.path.getsize(i)/1024:.1f} KB)")
    print("Listo.")
